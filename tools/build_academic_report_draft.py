from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "report" / "SANA学术报告初稿.docx"

TITLE = "基于 SANA 的高效文生图复现与 Prompt 复杂度自适应采样策略研究"


def set_run_font(run, size: float | None = None, east_asia: str = "宋体", ascii_font: str = "Times New Roman", bold: bool | None = None):
    run.font.name = ascii_font
    run._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia)
    run._element.rPr.rFonts.set(qn("w:ascii"), ascii_font)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), ascii_font)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold


def set_paragraph_format(paragraph, first_line: bool = True, line_pt: float = 20.0, before: float = 0, after: float = 0):
    pf = paragraph.paragraph_format
    pf.line_spacing = Pt(line_pt)
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    if first_line:
        pf.first_line_indent = Pt(24)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


def add_page_field(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    r = OxmlElement("w:r")
    t = OxmlElement("w:t")
    t.text = "1"
    r.append(t)
    fld.append(r)
    run._r.append(fld)
    set_run_font(run, 9)


def set_page_number_format(section, fmt: str = "decimal", start: int = 1):
    sect_pr = section._sectPr
    pg_num = sect_pr.find(qn("w:pgNumType"))
    if pg_num is None:
        pg_num = OxmlElement("w:pgNumType")
        sect_pr.append(pg_num)
    pg_num.set(qn("w:start"), str(start))
    pg_num.set(qn("w:fmt"), fmt)


def set_cell_text(cell, text: str, bold: bool = False, align=WD_ALIGN_PARAGRAPH.CENTER):
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.line_spacing = Pt(15)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    set_run_font(run, 10.5, bold=bold)


def shade_cell(cell, fill: str = "D9EAF7"):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float] | None = None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    for i, h in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], h, bold=True)
        shade_cell(table.rows[0].cells[i], "E7EEF7")
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], value, align=WD_ALIGN_PARAGRAPH.CENTER)
    if widths:
        for row in table.rows:
            for cell, w in zip(row.cells, widths):
                cell.width = Cm(w)
    doc.add_paragraph()
    return table


def add_paragraph(doc: Document, text: str):
    p = doc.add_paragraph()
    set_paragraph_format(p, first_line=True)
    run = p.add_run(text)
    set_run_font(run, 12)
    return p


def add_center_paragraph(doc: Document, text: str, size: float, bold: bool = False, east_asia: str = "黑体", before: float = 0, after: float = 0):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.line_spacing = Pt(20)
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    r = p.add_run(text)
    set_run_font(r, size, east_asia=east_asia, bold=bold)
    return p


def add_chapter(doc: Document, title: str):
    doc.add_page_break()
    p = doc.add_paragraph()
    p.style = doc.styles["Heading 1"]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(24)
    p.paragraph_format.space_after = Pt(18)
    p.paragraph_format.line_spacing = Pt(20)
    r = p.add_run(title)
    set_run_font(r, 16, east_asia="黑体", bold=True)
    return p


def add_section_heading(doc: Document, title: str, level: int = 2):
    p = doc.add_paragraph()
    p.style = doc.styles["Heading 2" if level == 2 else "Heading 3"]
    p.paragraph_format.space_before = Pt(24 if level == 2 else 12)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = Pt(20)
    p.paragraph_format.first_line_indent = None
    r = p.add_run(title)
    set_run_font(r, 14 if level == 2 else 13, east_asia="黑体" if level == 2 else "宋体", bold=False)
    return p


def add_figure(doc: Document, rel_path: str, caption: str, width_in: float = 6.2):
    path = ROOT / rel_path
    if not path.exists():
        add_paragraph(doc, f"（图像文件缺失：{rel_path}）")
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(path), width=Inches(width_in))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_before = Pt(6)
    cap.paragraph_format.space_after = Pt(12)
    cap.paragraph_format.line_spacing = Pt(16)
    r = cap.add_run(caption)
    set_run_font(r, 11, east_asia="宋体", bold=True)


def configure_styles(doc: Document):
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(12)

    for style_name in ["Heading 1", "Heading 2", "Heading 3"]:
        style = doc.styles[style_name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体" if style_name != "Heading 3" else "宋体")


def configure_section(section, header: bool = True):
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)
    section.header_distance = Cm(0.98)
    section.footer_distance = Cm(1.75)
    if header:
        hp = section.header.paragraphs[0]
        hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        hp.text = ""
        r = hp.add_run(TITLE)
        set_run_font(r, 9)
    fp = section.footer.paragraphs[0]
    fp.text = ""
    add_page_field(fp)


def build():
    doc = Document()
    configure_styles(doc)

    first = doc.sections[0]
    configure_section(first, header=False)
    set_page_number_format(first, fmt="upperRoman", start=1)

    add_center_paragraph(doc, TITLE, 20, bold=True, east_asia="黑体", before=12, after=18)
    add_center_paragraph(doc, "摘要", 16, east_asia="黑体", before=12, after=12)

    abstract_paragraphs = [
        "高分辨率文生图模型在视觉内容生成中表现出强大的语义表达能力，但扩散式生成过程通常需要多步迭代采样，推理成本仍然是实际应用中的重要约束。SANA 通过深度压缩自编码器、Linear Diffusion Transformer、decoder-only text encoder 与 Flow-DPM-Solver 等设计显著提高了高分辨率图像合成效率，但在实际推理中仍常采用固定采样步数和固定 guidance scale。不同文本提示词的对象数量、属性约束、空间关系和风格要求存在明显差异，固定采样策略可能在简单 prompt 上造成计算浪费，也可能在复杂 prompt 上出现预算不足。",
        "本项目在单张 NVIDIA GeForce RTX 5080 Laptop GPU 上复现了 SANA-0.6B 的 diffusers 推理流程，并构建了由 10-word、30-word 和 50-word 三组 prompt 组成的 30 条受控 benchmark。针对固定采样策略的不足，本文提出 Prompt-Complexity Adaptive Sampling（PCAS），在不修改 SANA 主模型参数的前提下，根据 prompt 复杂度动态分配采样步数和 guidance scale。除基于长度和结构特征的 Rule-PCAS 外，本文还实现了 DeepSeek-assisted PCAS，用于探索大语言模型语义复杂度判断在推理调度中的作用。",
        "实验表明，原始 Rule-PCAS 能够在短 prompt 上显著节省推理时间，但在均衡 benchmark 中会被长 prompt 的额外步数部分抵消，整体仅比 Fixed-20 快约 1.9%。为解决这一问题，本文进一步提出预算约束下的 Balanced-PCAS，将 10/30/50-word prompt 的采样步数调整为 8/16/24，使平均步数从 20 降至 16，去除预热后的平均推理时间从 0.996s 降至 0.751s，整体加速约 24.6%，同时 CLIPScore 基本保持在相同水平。对于 LLM-based 复杂度估计，原始 DeepSeek-PCAS 因将 19/30 条 prompt 判为 high 而过于保守，平均时间比 Fixed-20 慢约 30.8%；DeepSeek-Balanced 在复用同一批语义标签的基础上将策略调整为 low/medium/high=8/16/22 steps，使平均推理时间降至 0.844s，相比 Fixed-20 快约 15.3%。",
        "质量评估显示，各固定步数方法与 PCAS 方法之间的 CLIPScore 差异较小，hard prompt 定性对比也未显示稳定的质量上限提升。因此，本文将 PCAS 定位为一种在保持自动文本对齐指标和视觉观感基本相近的前提下改善效率-质量权衡的自适应推理策略，而非显著提升图像质量的方法。作为扩展实验，本文还复现了 SANA DreamBooth LoRA 个性化流程，并使用个人耳机照片分析小样本 LoRA 的主体一致性问题。结果显示，小样本 LoRA 在该硬件条件下可行，但主体保持能力受到数据质量、caption 设计和 adapter scale 的明显影响，应作为探索性增强而非强定量结论。",
    ]
    for text in abstract_paragraphs:
        add_paragraph(doc, text)

    p = doc.add_paragraph()
    set_paragraph_format(p, first_line=False)
    r = p.add_run("关键词")
    set_run_font(r, 12, bold=True)
    r = p.add_run("：SANA；文生图；Prompt 复杂度；自适应采样；高效推理；CLIPScore；LoRA")
    set_run_font(r, 12)

    body_section = doc.add_section(WD_SECTION.NEW_PAGE)
    body_section.header.is_linked_to_previous = False
    body_section.footer.is_linked_to_previous = False
    configure_section(body_section, header=True)
    set_page_number_format(body_section, fmt="decimal", start=1)

    # Chapter 1
    p = doc.add_paragraph()
    p.style = doc.styles["Heading 1"]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(24)
    p.paragraph_format.space_after = Pt(18)
    p.paragraph_format.line_spacing = Pt(20)
    r = p.add_run("第一章 引言")
    set_run_font(r, 16, east_asia="黑体", bold=True)

    add_section_heading(doc, "1.1 高分辨率文生图的效率问题")
    add_paragraph(doc, "扩散模型和 Diffusion Transformer 近年来显著推动了文生图质量提升，但其核心推理过程仍依赖多步去噪或流匹配采样。对于高分辨率图像生成而言，latent token 数量、注意力计算以及文本条件注入都会增加推理开销。实际应用中，用户往往希望在图像质量可接受的前提下尽可能降低单张图像的生成时间，因此如何在不重新训练大型模型的条件下优化推理预算，是高效文生图系统中的重要问题。")
    add_paragraph(doc, "SANA 是一个面向高效高分辨率图像合成的代表性模型。它通过更高压缩率的 autoencoder 减少 latent 空间 token 数，通过线性注意力降低 DiT 在高分辨率场景中的复杂度，并使用高效采样器进一步压缩采样步骤。尽管 SANA 本身已经具有较好的推理效率，实际调用时仍常使用统一的采样步数和 guidance scale，这意味着所有 prompt 被分配相同计算预算。")

    add_section_heading(doc, "1.2 固定采样参数的不足")
    add_paragraph(doc, "不同 prompt 的生成难度并不相同。简单 prompt 往往只包含单一主体和少量属性，例如一只苹果、一辆汽车或一个杯子；中等 prompt 会同时包含主体、场景、光照和风格；复杂 prompt 则可能包含多主体、多空间关系、文字生成和密集构图。对这些 prompt 使用完全相同的采样步数，容易产生两类问题：简单 prompt 上的额外采样步数可能带来边际收益递减，复杂 prompt 上过低的采样预算则可能导致主体缺失、关系不稳定或局部细节不足。")
    add_paragraph(doc, "因此，本项目关注的核心问题是：能否在不改变 SANA 主模型参数的前提下，根据 prompt 复杂度动态调整推理参数，使简单 prompt 更快，同时让复杂 prompt 保持较稳定的文本对齐和视觉观感。与训练新模型相比，这一路线更符合课程项目的时间和硬件约束，也更接近实际部署中的 inference-time adaptive computation 问题。")

    add_section_heading(doc, "1.3 本文贡献")
    for item in [
        "复现 SANA-0.6B diffusers 推理流程，在单张 RTX 5080 Laptop GPU 上完成 512×512 文生图实验，并保存环境、配置、输出图像和 metadata。",
        "构建 30 条受控 prompt benchmark，按 10-word、30-word 和 50-word 分为三组，系统比较 Fixed-10、Fixed-20 和 Fixed-28 固定步数 baseline。",
        "提出并实现 Prompt-Complexity Adaptive Sampling，包括 rule-based PCAS 和 DeepSeek-assisted PCAS，用于根据 prompt 复杂度动态选择 steps 和 guidance scale。",
        "针对原始 PCAS 整体加速不足和 DeepSeek-PCAS 过于保守的问题，提出 Balanced-PCAS 与 DeepSeek-Balanced，在保持 CLIPScore 接近的同时显著降低平均推理时间。",
        "系统评估采样步数、guidance scale 和 prompt complexity 对速度与 CLIPScore 的影响，并通过 hard prompt 定性分析给出更稳健的质量解释。",
        "作为可选增强，复现 SANA DreamBooth LoRA 个性化流程，并围绕主体一致性不稳定问题设计 enhanced LoRA、clean-captioned LoRA 和 subject-focused 验证。",
    ]:
        add_paragraph(doc, item)

    # Chapter 2
    add_chapter(doc, "第二章 相关工作与方法")
    add_section_heading(doc, "2.1 SANA Backbone")
    add_paragraph(doc, "SANA 的目标是在高分辨率图像合成中降低计算成本。其一，Deep Compression Autoencoder 将图像压缩到更少的 latent tokens，使后续扩散过程不必在过大的 token 序列上运行；其二，Linear Diffusion Transformer 用 linear attention 替代标准 attention，以降低高分辨率场景中的注意力开销；其三，decoder-only text encoder 用于增强文本理解和复杂 prompt 表达能力；其四，Flow-DPM-Solver 进一步提高采样效率。本文不重新训练 SANA 主模型，而是使用官方 diffusers 格式的 SANA-0.6B / 512px checkpoint 作为 backbone。")
    add_paragraph(doc, "从本项目角度看，SANA 的高效 backbone 使单卡复现和系统化消融成为可能，而其文本理解与采样设计也为 prompt-aware 的推理调度提供了自然动机。若模型本身已经能够在较少步数下给出可用结果，则更有必要追问不同 prompt 是否真的需要相同数量的采样步骤。")

    add_section_heading(doc, "2.2 Prompt 复杂度估计")
    add_paragraph(doc, "本文将 prompt 复杂度视为影响采样预算分配的代理变量。理想情况下，复杂度可由文本长度、主体数量、属性数量、空间和动作关系数量、风格约束、是否包含文字生成需求等因素共同决定。为了保证一周内可实现和可复现，本文首先采用 rule-based 简化实现，将 prompt 主要按词数分为 short、medium 和 long 三类，并为不同组分配不同采样策略。")
    add_paragraph(doc, "除长度规则外，本文还实现了 DeepSeek-assisted 复杂度估计。该模块调用 DeepSeek 模型对 prompt 中对象数量、关系数量、风格约束和文字生成需求进行结构化分析，并输出 low、medium 或 high 三类语义复杂度标签。为避免重复 API 调用，所有 DeepSeek 结果被缓存到本地 JSON 文件。DeepSeek 模块的意义不在于替代所有规则，而在于观察语义级复杂度判断能否改善仅依赖词数的调度策略。")

    add_section_heading(doc, "2.3 自适应采样策略")
    add_paragraph(doc, "PCAS 的整体流程为：输入 prompt 后先进行复杂度估计，再由 adaptive policy 选择采样步数、guidance scale、图像尺寸等推理参数，然后调用 SANA pipeline 生成图像，最后使用耗时、显存、CLIPScore 和定性图进行评估。本项目实际主线聚焦于 512×512 分辨率下的 steps 与 guidance scale 调度，adaptive resolution 保留为后续工作。")
    add_table(
        doc,
        ["方法", "复杂度来源", "采样策略", "定位"],
        [
            ["Fixed-20", "无", "所有 prompt 使用 20 steps / guidance 4.5", "主要固定步数 baseline"],
            ["Rule-PCAS", "词数分组", "10/30/50-word prompt 使用 10/20/28 steps", "展示 prompt-aware 计算重分配"],
            ["Balanced-PCAS", "词数分组", "10/30/50-word prompt 使用 8/16/24 steps", "预算约束下的效率主推版本"],
            ["DeepSeek-PCAS", "LLM 语义标签", "low/medium/high 使用 10/20/28 steps", "展示语义复杂度判断能力"],
            ["DeepSeek-Balanced", "LLM 语义标签", "low/medium/high 使用 8/16/22 steps", "预算约束下的语义调度"],
        ],
        widths=[3.2, 3.0, 6.0, 4.6],
    )

    add_section_heading(doc, "2.4 SANA-LoRA 个性化扩展")
    add_paragraph(doc, "在主线 PCAS 实验之外，本文复现 SANA DreamBooth LoRA 作为扩展实验。LoRA 通过在基础模型中插入低秩可训练适配器，以较小参数量完成个性化微调。本文使用 9 张个人耳机照片构建 DreamBooth 数据集，训练一个以 zzmearphone headphones 为触发词的 LoRA adapter，并在后续实验中比较原始 LoRA、增强 LoRA 和 clean-captioned LoRA 的主体一致性。该部分用于展示 SANA pipeline 的可扩展性和小样本个性化局限，不作为 PCAS 主线贡献。")

    # Chapter 3
    add_chapter(doc, "第三章 实验设置")
    add_section_heading(doc, "3.1 环境与硬件")
    add_paragraph(doc, "实验环境为 Windows 系统下的 Python 3.11 虚拟环境，核心依赖包括 PyTorch CUDA、diffusers、transformers、accelerate、safetensors、Pillow、pandas 和 pyyaml。GPU 为 NVIDIA GeForce RTX 5080 Laptop GPU，显存约 16GB。所有主要图像生成实验均基于 Efficient-Large-Model/Sana_600M_512px_diffusers，使用本地缓存模型并在 512×512 分辨率下运行。")
    add_table(
        doc,
        ["项目", "设置"],
        [
            ["模型", "Efficient-Large-Model/Sana_600M_512px_diffusers"],
            ["硬件", "NVIDIA GeForce RTX 5080 Laptop GPU, 15.92GB VRAM"],
            ["分辨率", "512×512"],
            ["随机种子", "42（主要 benchmark），242（LoRA subject-focused 验证）"],
            ["主要指标", "平均推理时间、去除预热后的平均时间、峰值显存、平均 steps、CLIPScore"],
            ["质量代理模型", "openai/clip-vit-base-patch32"],
        ],
        widths=[4.0, 12.0],
    )

    add_section_heading(doc, "3.2 Prompt Benchmark")
    add_paragraph(doc, "本文构建了 30 条受控 prompt benchmark，并按词数分为 10-word、30-word 和 50-word 三组，每组 10 条。短 prompt 多为单主体和简单场景，中等 prompt 包含较完整的环境、光照和风格描述，长 prompt 包含多主体、多关系、文字生成或密集构图。该设计使 Rule-PCAS 能够基于 prompt 长度进行确定性分组，也便于观察短、中、长 prompt 在固定采样和自适应采样下的速度差异。")
    add_figure(doc, "results/figures/day2_speed_summary_chart.png", "图3.1 固定采样步数 baseline 的平均推理时间对比", width_in=5.8)

    add_section_heading(doc, "3.3 Baseline 与评价指标")
    add_paragraph(doc, "固定步数 baseline 包括 Fixed-10、Fixed-20 和 Fixed-28，均使用 guidance scale=4.5，其中 Fixed-20 作为主要参考点。效率指标记录每张图像的推理时间、去除首张 warm-up 后的平均时间、峰值显存和平均采样步数。质量指标主要采用 CLIPScore，即 CLIP 图像特征与文本特征余弦相似度乘以 100。需要强调的是，CLIPScore 只能反映自动语义对齐，不能完全代表局部结构、主体完整性或人眼主观质量。")

    # Chapter 4
    add_chapter(doc, "第四章 实验结果")
    add_section_heading(doc, "4.1 固定步数 Baseline")
    add_paragraph(doc, "固定步数实验首先验证采样步数与推理时间之间的基本关系。Fixed-10 的去除预热平均时间为 0.525s，Fixed-20 为 0.996s，Fixed-28 为 1.335s。随着 steps 增加，推理时间近似上升；CLIPScore 也有轻微提升，但提升幅度较小。Fixed-28 的平均 CLIPScore 为 35.890，相比 Fixed-20 的 35.762 仅提升 0.128，却带来明显额外耗时。")
    add_table(
        doc,
        ["方法", "平均 steps", "去预热平均时间", "平均 CLIPScore"],
        [
            ["Fixed-10", "10.000", "0.525s", "35.689"],
            ["Fixed-20", "20.000", "0.996s", "35.762"],
            ["Fixed-28", "28.000", "1.335s", "35.890"],
        ],
        widths=[4.0, 3.5, 4.5, 4.0],
    )

    add_section_heading(doc, "4.2 Rule-PCAS 与 Balanced-PCAS")
    add_paragraph(doc, "原始 Rule-PCAS 根据 prompt 长度重新分配采样预算：10-word prompt 使用 10 steps，30-word prompt 使用 20 steps，50-word prompt 使用 28 steps。该策略在短 prompt 上节省明显，但长 prompt 使用比 Fixed-20 更多的采样步数，因此在 10/30/50 words 各 10 条的均衡 benchmark 上整体收益被抵消，去除预热平均时间只从 0.996s 降至 0.977s，约 1.9% 加速。")
    add_paragraph(doc, "为解决这一问题，本文提出效率优先的 Balanced-PCAS。它仍然保留 prompt-aware 的思想，但在平均计算预算上施加约束，将三组 prompt 的 steps 调整为 8/16/24。该策略使平均 steps 从 20 降至 16，整体平均时间降至 0.751s，相比 Fixed-20 加速约 24.6%。同时，Balanced-PCAS 的平均 CLIPScore 为 35.772，与 Fixed-20 的 35.762 基本一致。")
    add_table(
        doc,
        ["方法", "平均 steps", "去预热平均时间", "相对 Fixed-20 时间节省", "平均 CLIPScore"],
        [
            ["Fixed-20", "20.000", "0.996s", "0.0%", "35.762"],
            ["Rule-PCAS", "19.333", "0.977s", "1.9%", "35.750"],
            ["Balanced-PCAS", "16.000", "0.751s", "24.6%", "35.772"],
        ],
        widths=[3.5, 3.0, 3.8, 4.0, 3.4],
    )
    add_figure(doc, "results/figures/day3_pcas_balanced_speed_quality_tradeoff.png", "图4.1 Balanced-PCAS 与 Fixed-20 / Rule-PCAS 的速度-质量权衡", width_in=5.9)

    add_section_heading(doc, "4.3 DeepSeek-assisted PCAS")
    add_paragraph(doc, "DeepSeek-PCAS 的目标是用 LLM 估计 prompt 的语义复杂度，而不是仅依赖词数。在当前 30 条 benchmark 中，DeepSeek 将 5 条 prompt 判为 low，6 条判为 medium，19 条判为 high。原始策略将 high 映射到 28 steps，导致平均 steps 达到 23.4，去除预热平均时间为 1.303s，比 Fixed-20 慢约 30.8%。这说明 LLM 标签本身具备语义信息，但标签到采样预算的映射过于保守。")
    add_paragraph(doc, "DeepSeek-Balanced 复用同一批缓存标签，但将策略改为 low=8、medium=16、high=22 steps。这样既保留了语义复杂度排序，又避免 high prompt 全部升到 28 steps。最终 DeepSeek-Balanced 的平均 steps 为 18.467，平均时间为 0.844s，相比 Fixed-20 快约 15.3%，平均 CLIPScore 为 35.813，仍处于同一水平。")
    add_table(
        doc,
        ["方法", "Low/Medium/High", "平均 steps", "去预热平均时间", "相对 Fixed-20 时间节省", "平均 CLIPScore"],
        [
            ["Fixed-20", "-", "20.000", "0.996s", "0.0%", "35.762"],
            ["DeepSeek-PCAS", "5 / 6 / 19", "23.400", "1.303s", "-30.8%", "35.829"],
            ["DeepSeek-Balanced", "5 / 6 / 19", "18.467", "0.844s", "15.3%", "35.813"],
        ],
        widths=[3.4, 3.2, 2.8, 3.5, 3.5, 3.0],
    )
    add_figure(doc, "results/figures/day3_deepseek_balanced_speed_quality_tradeoff.png", "图4.2 DeepSeek-Balanced 将语义复杂度判断转化为预算约束调度", width_in=5.9)

    add_section_heading(doc, "4.4 质量评价与 Guidance Scale 消融")
    add_paragraph(doc, "在全部 30 条 prompt 上，各方法之间的 CLIPScore 差异很小。Fixed-10、Fixed-20、Fixed-28、Rule-PCAS、DeepSeek-PCAS 与 Balanced-PCAS 的平均 CLIPScore 均集中在 35.7 至 35.9 附近。因此，本文不将 PCAS 表述为显著提升图像质量的方法，而是将其解释为在保持自动文本对齐指标基本相近的同时改善计算预算分配。")
    add_paragraph(doc, "Guidance scale 消融固定采样步数为 20，并测试 1.5、3.5、4.5、5.5、6.5 和 8.5。结果显示 3.5-6.5 的中等区间较为稳定，平均 CLIPScore 范围仅约 0.130；guidance=1.5 明显降低文本对齐；guidance=8.5 给出最高 CLIPScore，但该现象更应解释为 CLIPScore 对高 guidance 的轻微偏好，而不是肉眼质量显著提升。")
    add_table(
        doc,
        ["Guidance scale", "去预热平均时间", "平均 CLIPScore"],
        [
            ["1.5", "0.897s", "34.705"],
            ["3.5", "1.190s", "35.856"],
            ["4.5", "0.996s", "35.762"],
            ["5.5", "1.124s", "35.848"],
            ["6.5", "1.088s", "35.892"],
            ["8.5", "0.887s", "36.027"],
        ],
        widths=[5.0, 5.0, 5.0],
    )
    add_figure(doc, "results/figures/day4_guidance_ablation_clipscore.png", "图4.3 Guidance scale 消融实验的 CLIPScore 变化", width_in=6.0)

    add_section_heading(doc, "4.5 Hard Prompt 定性分析")
    add_paragraph(doc, "为避免仅依赖平均 CLIPScore，本文从 benchmark 中选取 8 条复杂 prompt，覆盖复杂场景、多主体、多关系、文字生成和密集构图。比较方法包括 Fixed-20、Fixed-28、Rule-PCAS、Balanced-PCAS 和 DeepSeek-Balanced。Hard subset 的平均 CLIPScore 仍处于较窄区间，结合可视化网格观察，不同策略的差异主要体现在纹理、光照、边缘和局部小物体细节，而不是语义内容、主体完整性或构图关系的稳定改善。")
    add_table(
        doc,
        ["方法", "图像数", "平均 steps", "去预热平均时间", "平均 CLIPScore"],
        [
            ["Fixed-20", "8", "20.000", "0.978s", "36.522"],
            ["Fixed-28", "8", "28.000", "1.357s", "36.718"],
            ["Rule-PCAS", "8", "26.000", "1.225s", "36.748"],
            ["Balanced-PCAS", "8", "22.000", "0.953s", "36.526"],
            ["DeepSeek-Balanced", "8", "22.000", "0.955s", "36.440"],
        ],
        widths=[4.0, 2.2, 3.0, 4.0, 3.5],
    )
    add_figure(doc, "results/figures/day4_hard_prompt_qualitative_grid.png", "图4.4 Hard prompt 子集下不同采样策略的定性对比", width_in=5.1)

    add_section_heading(doc, "4.6 SANA-LoRA 个性化实验")
    add_paragraph(doc, "LoRA 扩展实验使用 9 张个人耳机照片作为 DreamBooth 数据集。原始 LoRA 使用 rank=8、alpha=8、200 steps；增强版 LoRA 使用 rank=16、alpha=16、500 steps，并将实例 prompt 明确为 black over-ear headphones with large oval ear cups and a padded headband；clean-captioned LoRA 则筛选 7 张更干净样本并为每张图添加 sidecar caption，训练 400 steps。")
    add_paragraph(doc, "初始验证显示，LoRA 能够成功训练和加载，但 scale=1.0 与 Base 差异较小，scale=2.0 在部分产品照 prompt 上更接近黑色头戴耳机，却没有在 CLIP 参考相似度上稳定超过 Base。为避免遮挡场景影响判断，本文进一步构造 8 条 subject-focused prompt，要求耳机居中且无遮挡。结果显示，Enhanced LoRA x1.5 在主体描述 CLIP 和 prompt CLIPScore 上最好，而 Clean-caption x1.25 是更保守的稳定方案；过高 scale 容易导致耳机结构坍塌或过度简化。")
    add_table(
        doc,
        ["方法", "Ref sim", "Subject CLIP", "Prompt CLIPScore", "简要结论"],
        [
            ["Base", "83.659", "28.487", "31.525", "无个性化 adapter"],
            ["Original LoRA x2", "83.839", "29.141", "32.620", "最接近参考中心，但提升很小"],
            ["Enhanced LoRA x1.5", "81.560", "30.053", "34.040", "自动主体描述与 prompt 对齐最佳"],
            ["Clean-caption x1.25", "83.281", "28.982", "33.084", "更保守、更稳定的 clean-data compromise"],
        ],
        widths=[4.0, 2.5, 3.0, 3.5, 5.0],
    )
    add_figure(doc, "results/figures/day5_lora_subject_consistency_grid.png", "图4.5 Day5 LoRA 主体一致性验证的多方法对比", width_in=5.6)

    # Chapter 5
    add_chapter(doc, "第五章 讨论")
    add_section_heading(doc, "5.1 PCAS 的主要价值")
    add_paragraph(doc, "本文实验表明，PCAS 的价值更准确地说是 adaptive computation，而不是显著提高生成质量。原始 Rule-PCAS 能说明 prompt-aware sampling 的直觉，即简单 prompt 少给 steps、复杂 prompt 多给 steps；但在均衡 benchmark 上，该策略会受到长 prompt 额外预算的抵消。Balanced-PCAS 通过将所有组的平均预算控制在 Fixed-20 以下，使自适应策略真正转化为整体加速。")
    add_paragraph(doc, "DeepSeek-Balanced 的结果进一步说明，LLM 复杂度标签并不能直接保证效率收益。语义复杂度判断本身有价值，但必须经过预算约束 policy 的校准，否则容易变成质量保守策略。对于实际部署而言，复杂度估计与动作策略应分开设计：前者负责排序与分类，后者负责在延迟、成本和质量之间进行约束优化。")

    add_section_heading(doc, "5.2 自动指标与定性评价的局限")
    add_paragraph(doc, "CLIPScore 可以快速衡量图像与 prompt 的语义相似度，但它对局部对象结构、复杂关系、文字细节和人眼审美并不敏感。本项目中不同策略的 CLIPScore 差异普遍较小，若直接宣称某一方法显著提升质量，会超过当前证据支持范围。更稳妥的解释是：Balanced-PCAS 在明显降低平均推理时间的同时，没有在 CLIP 自动对齐指标和定性观感上造成明显退化。")
    add_paragraph(doc, "Hard prompt 分析补充了这一结论。对复杂场景而言，不同采样策略确实会带来局部纹理和光照差异，但多数结果没有稳定的语义优劣。未来如果要更严格评价质量，应加入人工评分、VQA-based object/relationship 检测、ImageReward 或 PickScore 等更多指标，并在更大 prompt 分布上进行统计检验。")

    add_section_heading(doc, "5.3 LoRA 个性化的局限")
    add_paragraph(doc, "Day5 的 LoRA 实验证明，在单张 RTX 5080 Laptop GPU 上复现 SANA-LoRA DreamBooth 小样本个性化流程是可行的；但该实验也显示，小样本主体一致性依赖训练图质量、caption 精细度和 adapter scale。增强 LoRA x1.5 在 subject-focused prompt 上给出最好的自动主体描述匹配，但参考图相似度下降，说明它更像强化了“黑色头戴耳机”语义，而不一定稳定复制具体产品身份。Clean-caption x1.25 则说明筛选数据和逐图 caption 可以得到更保守的稳定结果。")
    add_paragraph(doc, "因此，LoRA 部分应作为探索性扩展呈现，而不应写成强结论。若未来目标是更清晰的产品身份复制，应补充 10-15 张干净产品图，覆盖正面、侧面、45 度、平放、立放、折叠、耳罩特写和头梁特写，并结合更精细的 identity-aware 评价指标。")

    add_section_heading(doc, "5.4 研究局限与未来工作")
    for item in [
        "Benchmark 规模仍较小，仅包含 30 条受控 prompt，尚未覆盖真实用户 prompt 分布。",
        "当前主线实验固定在 512×512 分辨率，原始大纲中设想的 adaptive resolution 尚未系统验证。",
        "质量评价主要依赖 CLIPScore 和定性图，缺少人工评分、VQA-based 关系检测和审美偏好指标。",
        "DeepSeek-assisted PCAS 只测试了一批缓存标签，尚未评估不同 LLM、不同 prompt 分布和 API 成本对策略的影响。",
        "LoRA 个性化数据量较少，主体一致性仍不稳定，当前结果只能支持流程可行性和局限分析。",
        "受硬件和时间限制，本文未全面测试 SANA 1.6B、更高分辨率或更强 backbone 下的迁移效果。",
    ]:
        add_paragraph(doc, item)

    # Chapter 6
    add_chapter(doc, "第六章 结论")
    add_paragraph(doc, "本文围绕高效文生图模型 SANA 展开，完成了 SANA-0.6B diffusers 推理流程复现、受控 prompt benchmark 构建、固定步数 baseline 对比、Prompt-Complexity Adaptive Sampling 设计与系统评估。实验结果表明，固定采样步数存在明显的计算预算分配不均问题，而 prompt-aware 的自适应采样可以在不修改主模型参数的情况下改善推理效率。")
    add_paragraph(doc, "具体而言，原始 Rule-PCAS 在短 prompt 上节省明显，但均衡 benchmark 中整体收益有限；Balanced-PCAS 通过 8/16/24 steps 的预算约束策略，将平均推理时间从 Fixed-20 的 0.996s 降至 0.751s，实现约 24.6% 加速，同时 CLIPScore 基本保持不变。DeepSeek-PCAS 证明了 LLM 语义复杂度判断的可行性，但原始策略过于保守；DeepSeek-Balanced 通过更温和的 low/medium/high 采样策略，将平均时间降至 0.844s，相比 Fixed-20 快约 15.3%。")
    add_paragraph(doc, "质量分析显示，各方法之间的 CLIPScore 差异较小，hard prompt 定性结果也未支持显著质量上限提升。因此，本文将 PCAS 定位为一种保持自动文本对齐指标并改善效率-质量权衡的推理调度策略，而不是显著提升图像质量的方法。最后，SANA-LoRA DreamBooth 扩展实验说明小样本个性化流程可行，但主体一致性仍受数据、caption 和 adapter scale 影响明显。整体而言，本项目完成了经典模型复现与轻量推理策略创新的课程目标，并为后续更大规模 prompt 分布、更多质量指标和 adaptive resolution 研究提供了基础。")

    # References
    add_chapter(doc, "参考文献")
    refs = [
        "Xie, E. et al. SANA: Efficient High-Resolution Image Synthesis with Linear Diffusion Transformers. arXiv:2410.10629, 2024.",
        "NVlabs. Sana official repository. https://github.com/NVlabs/Sana.",
        "Xie, E. et al. SANA 1.5: Efficient Scaling of Training-Time and Inference-Time Compute in Linear Diffusion Transformer. arXiv:2501.18427, 2025.",
        "Ho, J., Jain, A. & Abbeel, P. Denoising Diffusion Probabilistic Models. NeurIPS, 2020.",
        "Peebles, W. & Xie, S. Scalable Diffusion Models with Transformers. ICCV, 2023.",
        "Rombach, R. et al. High-Resolution Image Synthesis with Latent Diffusion Models. CVPR, 2022.",
        "Radford, A. et al. Learning Transferable Visual Models From Natural Language Supervision. ICML, 2021.",
        "Hu, E. J. et al. LoRA: Low-Rank Adaptation of Large Language Models. ICLR, 2022.",
        "Ruiz, N. et al. DreamBooth: Fine Tuning Text-to-Image Diffusion Models for Subject-Driven Generation. CVPR, 2023.",
        "Hessel, J. et al. CLIPScore: A Reference-free Evaluation Metric for Image Captioning. EMNLP, 2021.",
    ]
    for i, ref in enumerate(refs, 1):
        p = doc.add_paragraph()
        p.style = doc.styles["Normal"]
        p.paragraph_format.left_indent = Pt(18)
        p.paragraph_format.first_line_indent = Pt(-18)
        p.paragraph_format.line_spacing = Pt(18)
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(f"[{i}] {ref}")
        set_run_font(r, 11)

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
